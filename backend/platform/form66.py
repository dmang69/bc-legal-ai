"""Supreme Court Civil Rules Form 66 — Petition scaffold (DOCX).

This is a **drafting scaffold**, not a court-approved fillable form.
It maps petition outline content into Form 66 section headings used in BC
practice. Supervising counsel must complete registry requirements, style of
cause, and verify all authorities on BC Laws / CanLII before filing.

Form 66 = petition commencing a proceeding (including judicial review).
Form 67 = response to petition (not generated here).
"""

from __future__ import annotations

import io
from dataclasses import dataclass, field
from datetime import datetime, timezone
from typing import Any, Optional

from backend.identity import UserInfo


# Official form identity (locked design correction)
FORM_NUMBER = "66"
FORM_TITLE = "Petition"
RULE_REF = "Supreme Court Civil Rules, Rule 2-1 / Form 66"
JRPA_REF = "Judicial Review Procedure Act, RSBC 1996, c 241 (verify on BC Laws)"


@dataclass
class Form66Parties:
    petitioner: str = "[PETITIONER — complete legal name]"
    respondents: list[str] = field(
        default_factory=lambda: [
            "[RESPONDENT — e.g. Residential Tenancy Branch / Attorney General of BC as required]",
        ]
    )
    registry: str = "[Registry — e.g. Vancouver]"
    file_number: str = "[Court file number — leave blank if commencing]"


@dataclass
class Form66ScaffoldResult:
    ok: bool
    form_number: str = FORM_NUMBER
    court_ready: bool = False
    filename: str = ""
    docx_bytes: bytes = field(default=b"", repr=False)
    outline_status: str = ""
    warnings: list[str] = field(default_factory=list)
    error: str = ""
    sections: list[str] = field(default_factory=list)

    def to_dict(self) -> dict[str, Any]:
        return {
            "ok": self.ok,
            "form_number": self.form_number,
            "form_title": FORM_TITLE,
            "rule_ref": RULE_REF,
            "court_ready": self.court_ready,
            "filename": self.filename,
            "outline_status": self.outline_status,
            "warnings": self.warnings,
            "error": self.error,
            "sections": self.sections,
            "bytes": len(self.docx_bytes),
        }


def _default_warnings() -> list[str]:
    return [
        "NOT LEGAL ADVICE. Scaffold only — not a registry-ready Form 66.",
        "Verify JRPA / ATA limitation (e.g. ATA s.57) and Form 66 requirements before filing.",
        "Every legal proposition must be re-verified on BC Laws / CanLII with pinpoints.",
        "Privilege, consent, and supervising lawyer approval required before court-ready export.",
        "Form 66 starts a petition; Form 67 is the response — do not invert.",
    ]


def build_form66_docx(
    *,
    matter_id: str,
    user: UserInfo,
    outline: Optional[dict[str, Any]] = None,
    parties: Optional[Form66Parties] = None,
    matter_title: str = "",
) -> Form66ScaffoldResult:
    """Build a Form 66-structured DOCX from optional petition outline dict."""
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError as e:
        return Form66ScaffoldResult(
            ok=False,
            error=f"python-docx not installed: {e}",
            warnings=_default_warnings(),
        )

    parties = parties or Form66Parties()
    outline = outline or {}
    warnings = _default_warnings()
    sections: list[str] = []

    doc = Document()

    # Caption
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("IN THE SUPREME COURT OF BRITISH COLUMBIA")
    run.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f"Form {FORM_NUMBER} ({FORM_TITLE})").bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(RULE_REF).italic = True

    doc.add_paragraph(f"Registry: {parties.registry}")
    doc.add_paragraph(f"Court File No.: {parties.file_number}")
    doc.add_paragraph("")

    doc.add_paragraph("BETWEEN:").runs[0].bold = True
    doc.add_paragraph(parties.petitioner)
    doc.add_paragraph("PETITIONER").runs[0].bold = True
    doc.add_paragraph("")
    doc.add_paragraph("AND:").runs[0].bold = True
    for r in parties.respondents:
        doc.add_paragraph(r)
    doc.add_paragraph("RESPONDENT(S)").runs[0].bold = True

    # Meta
    doc.add_heading("Matter metadata (system)", level=2)
    sections.append("matter_metadata")
    doc.add_paragraph(f"Matter ID: {matter_id}")
    if matter_title:
        doc.add_paragraph(f"Matter title: {matter_title}")
    doc.add_paragraph(f"Draft prepared by user: {user.email} ({user.user_id})")
    doc.add_paragraph(f"Generated (UTC): {datetime.now(timezone.utc).isoformat()}")
    doc.add_paragraph(f"JR / statute route (verify): {JRPA_REF}")

    # Part 1 — Orders sought
    doc.add_heading("Part 1: ORDERS SOUGHT", level=2)
    sections.append("orders_sought")
    orders = outline.get("orders_sought") or outline.get("relief") or []
    if isinstance(orders, str):
        orders = [orders]
    if not orders:
        orders = [
            "[Order in the nature of certiorari / set aside RTB decision — particularize]",
            "[Order remitting the matter for rehearing — if sought]",
            "[Costs — if sought]",
            "[Such further and other relief as this Honourable Court deems just]",
        ]
        warnings.append("Orders sought are placeholders — particularize from record.")
    for o in orders:
        doc.add_paragraph(str(o), style="List Number")

    # Part 2 — Factual basis
    doc.add_heading("Part 2: FACTUAL BASIS", level=2)
    sections.append("factual_basis")
    facts = outline.get("material_facts") or outline.get("facts") or []
    if isinstance(facts, str):
        facts = [facts]
    if not facts:
        doc.add_paragraph(
            "[Material facts — FACT only, with record pinpoints (page/paragraph/timestamp). "
            "Do not mix legal argument here.]"
        )
        warnings.append("Material facts not supplied — insert HUMAN_CONFIRMED facts only.")
    else:
        for f in facts:
            doc.add_paragraph(str(f), style="List Number")

    # Part 3 — Legal basis / grounds
    doc.add_heading("Part 3: LEGAL BASIS", level=2)
    sections.append("legal_basis")
    grounds = outline.get("grounds") or []
    if grounds:
        for g in grounds:
            if not isinstance(g, dict):
                doc.add_paragraph(str(g), style="List Bullet")
                continue
            title = g.get("title") or g.get("ground_id") or "Ground"
            std = g.get("standard") or ""
            heading = f"Ground {g.get('ground_id', '')}: {title}"
            if std:
                heading += f" ({std})"
            doc.add_heading(heading.strip(), level=3)
            for sg in g.get("sub_grounds") or []:
                if not isinstance(sg, dict):
                    doc.add_paragraph(str(sg), style="List Bullet")
                    continue
                desc = sg.get("description") or ""
                doc.add_paragraph(f"{sg.get('sub_id', '')}: {desc}".strip(": "))
                for c in sg.get("cites") or []:
                    if isinstance(c, dict):
                        kind = c.get("kind") or "cite"
                        label = c.get("citation_short") or c.get("label") or ""
                        status = c.get("verification_status") or "UNVERIFIED"
                        doc.add_paragraph(
                            f"  [{kind}] {label} — status: {status}",
                            style="List Bullet",
                        )
                    else:
                        doc.add_paragraph(f"  {c}", style="List Bullet")
                if sg.get("notes"):
                    doc.add_paragraph(f"  Note: {sg['notes']}")
    else:
        doc.add_paragraph(
            "[Plead grounds with standard of review (e.g. patent unreasonableness / "
            "procedural fairness) and authorities — all UNVERIFIED until citation gate.]"
        )
        # Pull related title from outline
        if outline.get("title"):
            doc.add_paragraph(f"Outline title: {outline.get('title')}")
        if outline.get("statute_route"):
            doc.add_paragraph(f"Statute route: {outline.get('statute_route')}")

    # Part 4 — Material to be relied on
    doc.add_heading("Part 4: MATERIAL TO BE RELIED ON", level=2)
    sections.append("material_to_be_relied_on")
    materials = outline.get("materials") or outline.get("record") or []
    if isinstance(materials, str):
        materials = [materials]
    if not materials:
        materials = [
            "Affidavit of [deponent] (Form 109) — to be sworn/affirmed",
            "Certified copy of RTB decision under review (or best available)",
            "Transcript / recording extracts with timestamps [if any]",
            "Book of Authorities (tabs) — verified authorities only",
        ]
    for m in materials:
        doc.add_paragraph(str(m), style="List Bullet")

    # Footer disclaimers
    doc.add_heading("System disclaimers (do not file as pleading text)", level=2)
    sections.append("disclaimers")
    for w in warnings:
        doc.add_paragraph(w, style="List Bullet")
    doc.add_paragraph(
        f"court_ready: false | outline_status: {outline.get('status', 'scaffold')} | "
        f"form: {FORM_NUMBER}"
    )

    buf = io.BytesIO()
    doc.save(buf)
    fname = f"Form66_Petition_scaffold_{matter_id}.docx"
    return Form66ScaffoldResult(
        ok=True,
        court_ready=False,
        filename=fname,
        docx_bytes=buf.getvalue(),
        outline_status=str(outline.get("status") or "scaffold"),
        warnings=warnings,
        sections=sections,
    )


def form66_from_matter(
    *,
    user: UserInfo,
    matter_id: str,
    matter_title: str = "",
    client_label: str = "",
) -> Form66ScaffoldResult:
    """Load drafting outline for matter and render Form 66 DOCX scaffold."""
    from backend.platform import drafting as drafting_mod

    outline = drafting_mod.petition_outline(user, matter_id)
    parties = Form66Parties(
        petitioner=client_label or "[PETITIONER — complete legal name]",
    )
    return build_form66_docx(
        matter_id=matter_id,
        user=user,
        outline=outline if isinstance(outline, dict) else {},
        parties=parties,
        matter_title=matter_title,
    )
