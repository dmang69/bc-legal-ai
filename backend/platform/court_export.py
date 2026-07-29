"""Court-ready package export — DOCX binder only after approved export manifest.

Never sets court_ready without an APPROVED manifest. Public demo blocked upstream.
"""

from __future__ import annotations

import io
import json
import zipfile
from dataclasses import dataclass, field
from datetime import datetime, timezone
from typing import Any, Optional

from backend.db import get_connection, init_db
from backend.identity import AuthError, UserInfo, get_identity_service


@dataclass
class CourtPackageResult:
    ok: bool
    matter_id: str
    manifest_id: str
    court_ready: bool
    filename: str = ""
    media_type: str = "application/zip"
    blockers: list[str] = field(default_factory=list)
    error: str = ""
    package_bytes: bytes = field(default=b"", repr=False)

    def to_dict(self) -> dict[str, Any]:
        return {
            "ok": self.ok,
            "matter_id": self.matter_id,
            "manifest_id": self.manifest_id,
            "court_ready": self.court_ready,
            "filename": self.filename,
            "media_type": self.media_type,
            "blockers": self.blockers,
            "error": self.error,
            "bytes": len(self.package_bytes),
        }


def _load_manifest(matter_id: str, manifest_id: str) -> Optional[dict[str, Any]]:
    with get_connection() as conn:
        row = conn.execute(
            """
            SELECT * FROM export_manifests
            WHERE matter_id = ? AND manifest_id = ?
            """,
            (matter_id, manifest_id),
        ).fetchone()
    if not row:
        return None
    item = dict(row)
    item["document_ids"] = json.loads(item.get("document_ids_json") or "[]")
    item["citation_ids"] = json.loads(item.get("citation_ids_json") or "[]")
    item["blockers"] = json.loads(item.get("blockers_json") or "[]")
    item["approvals"] = json.loads(item.get("approvals_json") or "{}")
    item["court_ready"] = bool(item.get("court_ready"))
    return item


def _build_docx_summary(
    *,
    matter_id: str,
    manifest: dict[str, Any],
    user: UserInfo,
) -> bytes:
    try:
        from docx import Document
    except ImportError as e:
        raise RuntimeError("python-docx not installed") from e

    doc = Document()
    doc.add_heading("BC Legal AI Associate — Export Package Summary", level=1)
    doc.add_paragraph(
        "NOT A LAWYER. NOT LEGAL ADVICE. Supervising lawyer must verify all content before filing."
    )
    doc.add_paragraph(f"Matter ID: {matter_id}")
    doc.add_paragraph(f"Manifest ID: {manifest['manifest_id']}")
    doc.add_paragraph(f"Status: {manifest.get('status')}")
    doc.add_paragraph(f"Court-ready flag: {manifest.get('court_ready')}")
    doc.add_paragraph(f"Exported by: {user.email} ({user.user_id})")
    doc.add_paragraph(f"Exported at (UTC): {datetime.now(timezone.utc).isoformat()}")
    doc.add_heading("Approvals", level=2)
    for k, v in (manifest.get("approvals") or {}).items():
        doc.add_paragraph(f"{k}: {v}", style="List Bullet")
    doc.add_heading("Documents", level=2)
    for d in manifest.get("document_ids") or []:
        doc.add_paragraph(str(d), style="List Bullet")
    doc.add_heading("Citation verification IDs", level=2)
    for c in manifest.get("citation_ids") or []:
        doc.add_paragraph(str(c), style="List Bullet")
    doc.add_heading("Blockers at package time", level=2)
    blockers = manifest.get("blockers") or []
    if not blockers:
        doc.add_paragraph("None recorded on approved manifest.")
    else:
        for b in blockers:
            doc.add_paragraph(str(b), style="List Bullet")
    doc.add_paragraph(
        "Official legislation must be re-verified on BC Laws (currency line) before filing."
    )
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_court_package(
    *,
    user: UserInfo,
    matter_id: str,
    manifest_id: str,
) -> CourtPackageResult:
    """ZIP containing DOCX summary + manifest JSON. Requires APPROVED manifest."""
    init_db()
    if not get_identity_service().can_access_matter(user, matter_id, min_level="read"):
        raise AuthError("Matter access denied")

    manifest = _load_manifest(matter_id, manifest_id)
    if not manifest:
        return CourtPackageResult(
            ok=False,
            matter_id=matter_id,
            manifest_id=manifest_id,
            court_ready=False,
            error="Manifest not found",
            blockers=["Manifest not found"],
        )

    blockers: list[str] = []
    if manifest.get("status") != "APPROVED":
        blockers.append(f"Manifest status is {manifest.get('status')}, not APPROVED")
    if not manifest.get("court_ready"):
        blockers.append("Manifest court_ready is false")
    if manifest.get("blockers"):
        blockers.extend(f"manifest:{b}" for b in manifest["blockers"])

    if blockers:
        return CourtPackageResult(
            ok=False,
            matter_id=matter_id,
            manifest_id=manifest_id,
            court_ready=False,
            blockers=blockers,
            error="Export package blocked",
        )

    try:
        docx_bytes = _build_docx_summary(matter_id=matter_id, manifest=manifest, user=user)
    except Exception as e:
        return CourtPackageResult(
            ok=False,
            matter_id=matter_id,
            manifest_id=manifest_id,
            court_ready=False,
            error=str(e),
            blockers=[str(e)],
        )

    # Form 66 petition scaffold (always court_ready=false inside; package may still ship)
    form66_name = ""
    form66_bytes = b""
    form66_meta: dict[str, Any] = {}
    try:
        from backend.platform.form66 import form66_from_matter
        from backend.platform.matters import get_matter_store

        matter = get_matter_store().get_matter(user, matter_id)
        f66 = form66_from_matter(
            user=user,
            matter_id=matter_id,
            matter_title=str(matter.get("title") or ""),
            client_label=str(matter.get("client_label") or ""),
        )
        if f66.ok:
            form66_name = f66.filename or "Form66_Petition_scaffold.docx"
            form66_bytes = f66.docx_bytes
            form66_meta = f66.to_dict()
    except Exception as e:
        form66_meta = {"ok": False, "error": str(e), "form_number": "66", "court_ready": False}

    zbuf = io.BytesIO()
    with zipfile.ZipFile(zbuf, "w", compression=zipfile.ZIP_DEFLATED) as zf:
        zf.writestr("export_summary.docx", docx_bytes)
        if form66_bytes:
            zf.writestr(f"forms/{form66_name}", form66_bytes)
            zf.writestr("forms/form66_meta.json", json.dumps(form66_meta, indent=2))
        zf.writestr(
            "manifest.json",
            json.dumps(
                {
                    "manifest_id": manifest_id,
                    "matter_id": matter_id,
                    "status": manifest.get("status"),
                    "court_ready": True,
                    "document_ids": manifest.get("document_ids"),
                    "citation_ids": manifest.get("citation_ids"),
                    "approvals": manifest.get("approvals"),
                    "form66_scaffold": form66_meta,
                    "disclaimer": (
                        "Not legal advice. Human lawyer must verify before filing. "
                        "Form 66 scaffold is not a completed pleading."
                    ),
                },
                indent=2,
            ),
        )
        zf.writestr(
            "README.txt",
            "BC Legal AI Associate court package\n"
            "NOT LEGAL ADVICE. Verify legislation on BC Laws.\n"
            "forms/ contains Form 66 petition SCAFFOLD (Form 66 ≠ Form 67).\n"
            f"manifest={manifest_id} matter={matter_id}\n",
        )

    fname = f"court_package_{matter_id}_{manifest_id}.zip"
    return CourtPackageResult(
        ok=True,
        matter_id=matter_id,
        manifest_id=manifest_id,
        court_ready=True,
        filename=fname,
        package_bytes=zbuf.getvalue(),
    )
